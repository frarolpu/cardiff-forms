"""
Diagnostic script to inspect Word document structure.
Helps understand table layout and shading before extraction.
"""

import json
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

def get_cell_shading(cell):
    """Extract background color/shading from a table cell."""
    try:
        tcPr = cell._element.get_or_add_tcPr()
        shd = tcPr.find(qn('w:shd'))
        if shd is not None:
            fill = shd.get(qn('w:fill'))
            return fill
        return None
    except:
        return None

def inspect_document():
    """Inspect Word document structure."""
    
    doc_path = Path("Section 4 Schedule of Reports Word Nuevo.docx")
    
    if not doc_path.exists():
        print(f"ERROR: Document not found: {doc_path}")
        return
    
    doc = Document(str(doc_path))
    
    print(f"Document contains {len(doc.tables)} table(s)\n")
    
    # Inspect first table
    if len(doc.tables) > 0:
        table = doc.tables[0]
        print(f"Table 1: {len(table.rows)} rows x {len(table.rows[0].cells)} columns")
        print(f"{'='*100}\n")
        
        # Show first 10 rows
        for row_idx, row in enumerate(table.rows[:10]):
            print(f"ROW {row_idx}:")
            for col_idx, cell in enumerate(row.cells):
                text = cell.text.strip()[:50]  # First 50 chars
                shading = get_cell_shading(cell)
                shade_info = f" [SHADE: {shading}]" if shading else ""
                print(f"  COL {col_idx}: '{text}'{shade_info}")
            print()
        
        print(f"\n{'='*100}")
        print("Scanning for form numbers (e.g., 2.1.1, 2.36.6)...")
        print(f"{'='*100}\n")
        
        import re
        form_count = 0
        
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                # Look for form numbers
                if re.search(r'\d+\.\d+\.\d+', text):
                    form_match = re.search(r'(\d+\.\d+\.\d+)', text)
                    if form_match:
                        form_num = form_match.group(1)
                        shading = get_cell_shading(cell)
                        shade_info = f" [SHADE: {shading}]" if shading else ""
                        print(f"ROW {row_idx}, COL {col_idx}: Found form {form_num} {shade_info}")
                        
                        # Show all cells in this row
                        print(f"  Full row content:")
                        for i, c in enumerate(row.cells):
                            c_text = c.text.strip()[:40]
                            c_shade = get_cell_shading(c)
                            shade_str = f" [{c_shade}]" if c_shade else ""
                            print(f"    [{i}]: {c_text}{shade_str}")
                        print()
                        
                        form_count += 1
                        if form_count >= 5:  # Show first 5 forms
                            break
            if form_count >= 5:
                break

if __name__ == "__main__":
    inspect_document()
